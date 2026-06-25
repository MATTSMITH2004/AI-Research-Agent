#!/usr/bin/env python3
"""Render a research-digest brief (markdown) to a styled .docx.

Usage: python3 scripts/md_to_docx.py <input.md> <output.docx>

Handles the subset of markdown the briefs use: # / ## / ### headings,
**bold**, *italic*, "- " bullets, "Source:" lines, --- rules, and bare URLs.
URLs are rendered as clickable blue hyperlinks (anchor text = the domain) so
every source is one click away and the page stays readable.
"""

import re
import sys
from urllib.parse import urlparse

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn

HYPERLINK_BLUE = "0563C1"
URL_RE = r"https?://[^\s)\]>;]+"
TOKEN = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|" + URL_RE + r")")


def _domain(url):
    netloc = urlparse(url).netloc
    return netloc[4:] if netloc.startswith("www.") else netloc


def add_hyperlink(paragraph, url, text, size=None):
    """Append a clickable, blue, underlined hyperlink run to the paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), HYPERLINK_BLUE)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))  # half-points
        rPr.append(sz)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _sized(run, size):
    if size is not None:
        run.font.size = Pt(size)
    return run


def add_inline(paragraph, text, size=None):
    """Add text, honoring **bold**, *italic*, and bare URLs (as blue links)."""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _sized(paragraph.add_run(part[2:-2]), size).bold = True
        elif part.startswith("*") and part.endswith("*"):
            _sized(paragraph.add_run(part[1:-1]), size).italic = True
        elif re.fullmatch(URL_RE, part):
            # peel trailing sentence punctuation back into plain text
            trail = ""
            while part and part[-1] in ".,":
                trail = part[-1] + trail
                part = part[:-1]
            add_hyperlink(paragraph, part, _domain(part), size=size)
            if trail:
                _sized(paragraph.add_run(trail), size)
        else:
            _sized(paragraph.add_run(part), size)


def main():
    src, dst = sys.argv[1], sys.argv[2]
    lines = open(src, encoding="utf-8").read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip() or line.strip() == "---":
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:].strip())
        elif line.startswith("Source:"):
            p = doc.add_paragraph()
            lbl = p.add_run("Source: ")
            lbl.bold = True
            lbl.font.size = Pt(9)
            # keep the "(what it supports)" notes; URLs become blue domain links
            add_inline(p, line[len("Source:"):].strip(), size=9)
        else:
            p = doc.add_paragraph()
            if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
                p.add_run(line[1:-1]).italic = True
            else:
                add_inline(p, line)

    doc.save(dst)
    print("wrote", dst)


if __name__ == "__main__":
    main()
